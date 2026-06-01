from pathlib import Path
from docx import Document
from typing import List, Optional
from app.models import CVProfile


def generate_cv_docx(
    profile: CVProfile,
    summary: str,
    skills: List[str],
    experience: List[str],
    education: Optional[str],
    output_path: Path,
) -> None:
    document = Document()
    document.add_heading(profile.name, level=0)

    contact = [profile.email]
    if profile.phone:
        contact.append(profile.phone)
    if profile.linkedin:
        contact.append(profile.linkedin)
    document.add_paragraph(" | ".join(contact))

    document.add_heading("Professional Summary", level=1)
    document.add_paragraph(summary)

    document.add_heading("Skills", level=1)
    skills_paragraph = document.add_paragraph()
    skills_paragraph.add_run(", ".join(skills))

    document.add_heading("Experience", level=1)
    for bullet in experience:
        p = document.add_paragraph(style="List Bullet")
        p.add_run(bullet)

    if education:
        document.add_heading("Education", level=1)
        document.add_paragraph(education)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
